# -*- coding: utf-8 -*-
"""
Created on Mon Jan 22 10:40:39 2018
@author: mot16

Measures the agreemet rates of physiscians with LIME explanations
"""
import pandas as pd
#import numpy as np
from docx import Document
from patient_pids import get_pids

tp_pids, tn_pids, fake_tp_pids, fake_tn_pids, all_tp_pids, all_tn_pids, pids = get_pids()

## read description documents
pt_descs = {}
pt_descs['Shyam'] = Document("../data/patient_descriptions_Shyam.docx")
pt_descs['Luca'] = Document("../data/patient_descriptions_Luca.docx")
pt_descs['Malar'] = Document("../data/patient_descriptions_Malar.docx")

## extract results from all description documents
results = pd.DataFrame(columns=['pt_no','pid',
'physician','tp_tn','fake_real','f1','f2','f3','f4','f5','f6'])

for name, doc in pt_descs.items():
    for i in range(len(doc.tables)):
        row_dict = {}
        row_dict['pt_no'] = i+1
        row_dict['pid'] = pids[i]
        row_dict['physician'] = name
        row_dict['tp_tn'] = 'tp' if pids[i] in all_tp_pids.values else 'tn'
        row_dict['fake_real'] = 'fake' if (pids[i] in fake_tp_pids.values) or (pids[i] in fake_tn_pids.values) else 'real'
        for j in range(1,7):
            row_dict['f'+str(j)] = doc.tables[i].rows[j].cells[1].text.strip()
        results = results.append(row_dict, ignore_index=True)

## save results
results.to_csv("../output/pt_descriptions_results.csv", index=False)


def replace(x):
    return 1 if x>=2 else 0
    
agreement_real = results.loc[results.fake_real == 'real',:].groupby([
                            'pt_no',
                             'pid',
                             'tp_tn',
                             'fake_real'],
                             sort=False)[results.columns.values[-6:]].apply(lambda x: x.sum()).applymap(replace).reset_index()

## agreement percentage in real cases
## Majority
agreement_real.iloc[:,-6:].sum().sum() / 180
## in true positives
agreement_real.loc[agreement_real.tp_tn == 'tp',:].iloc[:,-6:].sum().sum() / 90
## in true negatives
agreement_real.loc[agreement_real.tp_tn == 'tn',:].iloc[:,-6:].sum().sum() / 90


agreement_real_Shyam = results.loc[(results.fake_real == 'real') & (results.physician == 'Shyam')]
## real all
agreement_real_Shyam.iloc[:,-6:].sum().sum() / 180
## in true positives
agreement_real_Shyam.loc[agreement_real_Shyam.tp_tn == 'tp',:].iloc[:,-6:].sum().sum() / 90
## in true negatives
agreement_real_Shyam.loc[agreement_real_Shyam.tp_tn == 'tn',:].iloc[:,-6:].sum().sum() / 90


agreement_real_Luca = results.loc[(results.fake_real == 'real') & (results.physician == 'Luca')]
## real all
agreement_real_Luca.iloc[:,-6:].sum().sum()
agreement_real_Luca.iloc[:,-6:].sum().sum() / 180
## in true positives
agreement_real_Luca.loc[agreement_real_Luca.tp_tn == 'tp',:].iloc[:,-6:].sum().sum()
agreement_real_Luca.loc[agreement_real_Luca.tp_tn == 'tp',:].iloc[:,-6:].sum().sum() / 90
## in true negatives
agreement_real_Luca.loc[agreement_real_Luca.tp_tn == 'tn',:].iloc[:,-6:].sum().sum()
agreement_real_Luca.loc[agreement_real_Luca.tp_tn == 'tn',:].iloc[:,-6:].sum().sum() / 90


agreement_real_Malar = results.loc[(results.fake_real == 'real') & (results.physician == 'Malar')]
## real all
agreement_real_Malar.iloc[:,-6:].sum().sum()
agreement_real_Malar.iloc[:,-6:].sum().sum() / 180
## in true positives
agreement_real_Malar.loc[agreement_real_Malar.tp_tn == 'tp',:].iloc[:,-6:].sum().sum()
agreement_real_Malar.loc[agreement_real_Malar.tp_tn == 'tp',:].iloc[:,-6:].sum().sum() / 90
## in true negatives
agreement_real_Malar.loc[agreement_real_Malar.tp_tn == 'tn',:].iloc[:,-6:].sum().sum()
agreement_real_Malar.loc[agreement_real_Malar.tp_tn == 'tn',:].iloc[:,-6:].sum().sum() / 90



## fake explanations
agreement_fake = results.loc[results.fake_real == 'fake',:].groupby([
                            'pt_no',
                             'pid',
                             'tp_tn',
                             'fake_real'],
                             sort=False)[results.columns.values[-6:]].apply(lambda x: x.sum()).applymap(replace).reset_index()

## agreement percentage in fake cases
## Majority
agreement_fake.iloc[:,-6:].sum().sum() / (agreement_fake.shape[0]*6)
## in true positives
agreement_fake.loc[agreement_fake.tp_tn == 'tp',:].iloc[:,-6:].sum().sum() / (agreement_fake.shape[0]*6/2)
## in true negatives
agreement_fake.loc[agreement_fake.tp_tn == 'tn',:].iloc[:,-6:].sum().sum() / (agreement_fake.shape[0]*6/2)


agreement_fake_Shyam = results.loc[(results.fake_real == 'fake') & (results.physician == 'Shyam')]
## fake all
agreement_fake_Shyam.iloc[:,-6:].sum().sum() / 60
## in true positives
agreement_fake_Shyam.loc[agreement_fake_Shyam.tp_tn == 'tp',:].iloc[:,-6:].sum().sum() / 30
## in true negatives
agreement_fake_Shyam.loc[agreement_fake_Shyam.tp_tn == 'tn',:].iloc[:,-6:].sum().sum() / 30


agreement_fake_Luca = results.loc[(results.fake_real == 'fake') & (results.physician == 'Luca')]
## fake all
agreement_fake_Luca.iloc[:,-6:].sum().sum()
agreement_fake_Luca.iloc[:,-6:].sum().sum() / 60
## in true positives
agreement_fake_Luca.loc[agreement_fake_Luca.tp_tn == 'tp',:].iloc[:,-6:].sum().sum()
agreement_fake_Luca.loc[agreement_fake_Luca.tp_tn == 'tp',:].iloc[:,-6:].sum().sum() / 30
## in true negatives
agreement_fake_Luca.loc[agreement_fake_Luca.tp_tn == 'tn',:].iloc[:,-6:].sum().sum()
agreement_fake_Luca.loc[agreement_fake_Luca.tp_tn == 'tn',:].iloc[:,-6:].sum().sum() / 30


agreement_fake_Malar = results.loc[(results.fake_real == 'fake') & (results.physician == 'Malar')]
## fake all
agreement_fake_Malar.iloc[:,-6:].sum().sum()
agreement_fake_Malar.iloc[:,-6:].sum().sum() / 60
## in true positives
agreement_fake_Malar.loc[agreement_fake_Malar.tp_tn == 'tp',:].iloc[:,-6:].sum().sum()
agreement_fake_Malar.loc[agreement_fake_Malar.tp_tn == 'tp',:].iloc[:,-6:].sum().sum() / 30
## in true negatives
agreement_fake_Malar.loc[agreement_fake_Malar.tp_tn == 'tn',:].iloc[:,-6:].sum().sum()
agreement_fake_Malar.loc[agreement_fake_Malar.tp_tn == 'tn',:].iloc[:,-6:].sum().sum() / 30



import seaborn as sns
#import matplotlib.gridspec as gridspec
import matplotlib.pyplot as plt

agreement_all = results.groupby(['pt_no', 'pid', 'tp_tn','fake_real'],
                                sort=False)[results.columns.values[-6:]].apply(lambda x: x.sum()).applymap(replace).reset_index()
agreement_all['total'] = agreement_all.iloc[:,-6:].sum(axis=1)

#plt.figure(figsize=(1,100))

sns.set_style('whitegrid')
#gs = gridspec.GridSpec(1, 10)
#ax1 = plt.subplot(gs[0, 0:5])
#ax2 = plt.subplot(gs[0, 5:])
fig, axes = plt.subplots(nrows=1,ncols=2, figsize=(12, 4))
tp_ax = sns.barplot(x="pt_no",y="total",
                    data=agreement_all.loc[(agreement_all.fake_real == 'fake')&(agreement_all.tp_tn=='tp')],
                     palette=['green'], ax=axes[0])
tp_ax.set_ylim(0,5)
tp_ax.set(xlabel='patient#', ylabel='#Majority agreement (out of 6 features)',
          title='Fake true positive explanations')
tn_ax = sns.barplot(x="pt_no",y="total",
                    data=agreement_all.loc[(agreement_all.fake_real == 'fake')&(agreement_all.tp_tn=='tn')],
                     palette=['blue'], ax=axes[1])
#sns.despine(ax=ax2, bottom=True)
#tp_ax.set_ylim(0,5)
tn_ax.set(xlabel='patient#', ylabel='',
          title='Fake true negative explanations')
