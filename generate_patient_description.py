# -*- coding: utf-8 -*-
"""
Created on Wed Nov 22 18:42:07 2017
@author: mot16

Generates a Word document that contains patient descriptions and LIME explanations as plots
"""
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches
from explanation_plot import draw_exp_plot


def desc_value_cat(value, arg1,out1,arg2=None,out2=None):
    out=''
    if value==-9:
        out = '-'
    elif value==arg1:
        out = out1
    elif value==arg2:
        out=out2
    return out
    
def desc_value_cont(value):
    out = '-' if value==-9 else value
    return out

def raw2desc_val(base, base_varname, pid, desc_vars):
    val = base.loc[pid, base_varname]
    if val == -9:
        return '-'
    var_type = desc_vars['base_type'][desc_vars.base == base_varname].item()
    if var_type == 'int':
        return val
    elif var_type == 'cat':
        return 'yes' if val == 1 else 'no'
    else:
        if base_varname == '120.SMOKE':
            return 'non-smoker' if val == 0 else 'smoker'
        if base_varname == '25.PTHISP':
            return 'non-Hispanic' if val == 0 else 'Hispanic'
        if base_varname == '26.PTRACEA':
            return 'white' if val == 0 else 'non-white'
        if base_varname == '23.SEX':
            return 'male' if val == 1 else 'female'
        if base_varname == '206.CXRINF':
            return 'yes' if val >= 1 else 'no'
        if base_varname == '145.PULCLEAR':
            return 'congested' if val == 0 else 'clear'
        if base_varname == '117.CONFUSA':
            return 'yes' if val >= 1 else 'no' 
            
def explanation_info(exp_df, pid):
    prob = exp_df.loc[exp_df.pid == pid, 'model_proba'].item()
    exp = exp_df.loc[exp_df.pid == pid].iloc[0, 5:].dropna()
    exp = exp[np.argsort(np.abs(exp))[::-1]]

    return prob, exp

def isin(list_, item):
    return np.sum(np.isin(list_, item)) > 0


def pt_actual_label(pid, tp_pids, tn_pids, fake_tp_pids, fake_tn_pids):
    if isin(tp_pids, pid):
        return 'TP'
    elif isin(fake_tp_pids, pid):
        return 'FAKE TP'
    elif isin(tn_pids, pid):
        return 'TN'
    elif isin(fake_tn_pids, pid):
        return 'FAKE TN'

def unit(desc_vars, varname, value):
    desc_vars.unit[desc_vars.unit.isnull()] = ''
    if value == '-':
        return ''
    
    return desc_vars.loc[desc_vars.base == varname,'unit'].item()    
    
    
tp_exp_df = pd.read_csv("../output/true_tp_explanations.csv")
fake_tp_exp_df = pd.read_csv("../output/fake_tp_explanations.csv")
tn_exp_df = pd.read_csv("../output/true_tn_explanations.csv")
fake_tn_exp_df = pd.read_csv("../output/fake_tn_explanations.csv")

port_test = pd.read_csv('../data/port_test_new.csv',index_col=0)
base = pd.read_csv('../data/base.csv', index_col=0)
desc_vars = pd.read_csv("../data/desc_vars_2.csv")

document = Document()        

tp_pids = tp_exp_df.sort_values('model_proba', ascending=False).pid[0:15]
fake_tp_pids = fake_tp_exp_df.sort_values('model_proba', ascending=False).pid[0:5]

tn_pids = tn_exp_df.sort_values('model_proba', ascending=False).pid[0:15]
fake_tn_pids = fake_tn_exp_df.sort_values('model_proba', ascending=False).pid[0:5]

pids1 = tp_pids.append(fake_tp_pids, ignore_index=True)
pids2 = tn_pids.append(fake_tn_pids, ignore_index=True)
pids = pids1.append(pids2 , ignore_index=True)
np.random.seed(0)
np.random.shuffle(pids)

for pid in pids:
    s = 'Patient#' + str(list(pids).index(pid)+1)
    
#    s+= " (" + pt_actual_label(pid, tp_pids, 
#                               tn_pids, fake_tp_pids,
#                               fake_tn_pids) + ' EXPLANATION)'
    
    document.add_paragraph(s)

    s = 'Demographics:\n'
    val=raw2desc_val(base, '20.AGEPRES', pid, desc_vars)
    s += 'Age=' + str(val) + unit(desc_vars, '20.AGEPRES', val) + ', '
    val=raw2desc_val(base, '23.SEX', pid, desc_vars)
    s += 'Sex=' + val + ', '
    val=raw2desc_val(base, '26.PTRACEA', pid, desc_vars)
    s += 'Race=' + str(val) + ', '
    val=raw2desc_val(base, '25.PTHISP', pid, desc_vars)
    s += 'Ethnicity=' + str(val) + ', '
    val=raw2desc_val(base, '120.SMOKE', pid, desc_vars)
    s += 'Smoking status=' + str(val)
    document.add_paragraph(s)
    
    s = 'Past history:\n'
    val=raw2desc_val(base, '60.NOPNEPIS', pid, desc_vars)
    s += 'Number of prior episodes of pneumonia=' + str(val)
    document.add_paragraph(s)

    s = 'Comorbidities:\n'
    val=raw2desc_val(base, '74.CHFA', pid, desc_vars)
    s += 'Congestive heart failure=' + str(val) + ', '
    val=raw2desc_val(base, '79.CVDA', pid, desc_vars)
    s += 'Cerebrovascular disease=' + str(val) + ', '
    val=raw2desc_val(base, '77.LIVERDIA', pid, desc_vars)
    s += 'Liver disease=' + str(val) + ', '
    val=desc_value_cat(base.loc[pid,'71.ACTIVCAA'],0,'absent',1,'present')       
    val=raw2desc_val(base, '71.ACTIVCAA', pid, desc_vars)
    s += 'Cancer=' + str(val)
    document.add_paragraph(s)
    
    s = 'Symptoms:\n'
    val=raw2desc_val(base, '28.COUGHY', pid, desc_vars)
    s += 'Cough=' + str(val) + ', '
    val=raw2desc_val(base, '39.FEVERY', pid, desc_vars)
    s += 'Fever=' + str(val) + ', '
    val=raw2desc_val(base, '40.SWEATSY', pid, desc_vars)
    s += 'Sweating=' + str(val) + ', '
    val=raw2desc_val(base, '44.HEADACHY', pid, desc_vars)
    s += 'Headache=' + str(val)
    document.add_paragraph(s)
    
    s = 'Physical exam:\n'
    val=raw2desc_val(base, '117.CONFUSA', pid, desc_vars)
    s += 'Confusion=' + str(val) + ', '
    val=raw2desc_val(base, '145.PULCLEAR', pid, desc_vars)
    s+='Lungs status=' + str(val)
    document.add_paragraph(s)
    
    s='Laboratory results:\n'
    val=raw2desc_val(base, '133.PULSE', pid, desc_vars)
    s += 'HR=' + str(val) + unit(desc_vars, '133.PULSE', val) + '   '
    bpsys=raw2desc_val(base, '131.BPSYS', pid, desc_vars)
    bpdias=raw2desc_val(base, '132.BPDIAS', pid, desc_vars)
    s += 'BP=' + str(bpsys) + '/' + str(bpdias) + unit(desc_vars, '131.BPSYS', val) + '   '
    val=raw2desc_val(base, '134.RESPRATE', pid, desc_vars)
    s += 'RR=' + str(val) + unit(desc_vars, '134.RESPRATE', val) + '   '
    val=raw2desc_val(base, '135.TEMPC', pid, desc_vars)
    s += 'Temp=' + str(val) + unit(desc_vars, '135.TEMPC', val)
    s += '\n'
    val=raw2desc_val(base, '158.WBC', pid, desc_vars)
    s += 'WBC=' + str(val) + unit(desc_vars, '158.WBC', val) + '   '
    val=raw2desc_val(base, '162.HGB', pid, desc_vars)
    s += 'Hgb=' + str(val) + unit(desc_vars, '162.HGB', val) + '   '
    val=raw2desc_val(base, '161.HCT', pid, desc_vars)
    s += 'Hct=' + str(val) + unit(desc_vars, '161.HCT', val) + '   '
    val=raw2desc_val(base, '163.PLT', pid, desc_vars)
    s += 'Plt=' + str(val) +  unit(desc_vars, '163.PLT', val)
    s += '\n'
    val=raw2desc_val(base, '171.NA', pid, desc_vars)
    s += 'Na=' + str(val) + unit(desc_vars, '171.NA', val) + '   '
    val=raw2desc_val(base, '172.KP', pid, desc_vars)
    s += 'K=' + str(val) + unit(desc_vars, '172.KP', val) + '   '
    val=raw2desc_val(base, '173.HCO3', pid, desc_vars)
    s += 'HCO3=' + str(val) + unit(desc_vars,'173.HCO3', val) + '   '
    val=raw2desc_val(base, '174.BUN', pid, desc_vars)
    s += 'BUN=' + str(val) + unit(desc_vars, '174.BUN', val) + '   '
    val=raw2desc_val(base, '175.CR', pid, desc_vars)
    s += 'Cr=' + str(val) + unit(desc_vars, '175.CR', val) + '   '
    val=raw2desc_val(base, '170.GLU', pid, desc_vars)
    s += 'Glu=' + str(val) +  unit(desc_vars, '170.GLU', val)
    s+='\n'
    val=raw2desc_val(base, '184.BILIR', pid, desc_vars)
    s += 'Tot Bili=' + str(val) + unit(desc_vars, '184.BILIR', val) + '   '
    val=raw2desc_val(base, '182.SGOT', pid, desc_vars)
    s += 'SGOT/AST=' + str(val) + unit(desc_vars, '182.SGOT', val) + '   '
    val=raw2desc_val(base, '183.ALKPHOS', pid, desc_vars)
    s += 'Alk Phos=' + str(val) + unit(desc_vars, '183.ALKPHOS', val) + '   '
    val=raw2desc_val(base, '186.LDH', pid, desc_vars)
    s += 'LDH=' + str(val) + unit(desc_vars, '186.LDH', val)
    s+='\n'
    s += 'ABG '
    val=raw2desc_val(base, '196.PH', pid, desc_vars)
    s += 'pH=' + str(val) + unit(desc_vars, '196.PH', val) + '   '
    val=raw2desc_val(base, '197.PCO2', pid, desc_vars)
    s += 'pCO2=' + str(val) + unit(desc_vars, '197.PCO2', val) + '   '
    val=raw2desc_val(base, '198.PO2', pid, desc_vars)
    s += 'pO2=' + str(val) + unit(desc_vars, '198.PO2', val)
    s +='\n'
    val = desc_value_cont(base.loc[pid,'192.O2SAT'])
    val=raw2desc_val(base, '192.O2SAT', pid, desc_vars)
    s += 'O2 saturation=' + str(val) + unit(desc_vars, '192.O2SAT', val)
    document.add_paragraph(s)
    
    s = 'X-ray:\n'
    val=raw2desc_val(base, '206.CXRINF', pid, desc_vars)            
    s += 'Infiltrate=' + str(val) + ', '
    val=raw2desc_val(base, '212.CXREFF', pid, desc_vars)
    s+= 'Pleural effusion=' + str(val)
    document.add_paragraph(s)
    
    s = 'Outcome:\n'
    val=port_test.loc[pid,'217.DIREOUT']
    val='yes' if val == 1 else 'no'
    #raw2desc_val(base, '217.DIREOUT', pid, desc_vars)
    s+='Dire outcome=' + str(val)
    document.add_paragraph(s)
    
    ## extract prob and exp features
    prob=exp_= 0
    if isin(tp_pids, pid):
        prob, exp_ = explanation_info(tp_exp_df, pid)
    elif isin(fake_tp_pids, pid):
        prob, exp_ = explanation_info(fake_tp_exp_df, pid)
    elif isin(tn_pids, pid):
        prob, exp_ = explanation_info(tn_exp_df, pid)
        prob = 1- prob
    elif isin(fake_tn_pids, pid):
        prob, exp_ = explanation_info(fake_tn_exp_df, pid)
        prob = 1- prob
    ## pick 6 features
    if isin(tp_pids,pid) or isin(tn_pids,pid):
        exp_features = exp_.index[0:6]
    elif isin(fake_tp_pids,pid) or isin(fake_tn_pids,pid):
        exp_features = exp_.index[-6:]
    
    ## map feature names to original names and add unit    
    exp_varnames = [f.split('=')[0] for f in exp_features]
    temp_desc_df = desc_vars[desc_vars.port.isin(exp_varnames)]
    temp_desc_df['sort_exp'] = pd.Categorical(temp_desc_df.port, categories=exp_varnames, ordered=True)
    temp_desc_df.sort_values('sort_exp', inplace=True)
    temp_desc_df['unit'][temp_desc_df['unit'].isnull()] = ''
    
    feature_values = temp_desc_df['base'].apply(lambda x: raw2desc_val(base, x, pid, desc_vars)).values
    feature_values = ['%.1f' % e if type(e) == np.float64 else str(e) for e in feature_values]
    feature_values = feature_values + temp_desc_df['unit'].values
    
    labels = temp_desc_df['desc_name'].values + '=' + feature_values
    weights = exp_.values[0:6]
    exp_df = pd.DataFrame({'labels':labels, 'weights':weights})
    
    draw_exp_plot(exp_df, round(prob,2), "../output/exp_plot.png")
    
    
    s = 'Explanation:'
    document.add_paragraph(s)
    
    document.add_picture("../output/exp_plot.png", width=Inches(6))
    
    
    s= "In the table below in the last column, please state your agreement/disagreement with each feature. You may indicate agreement with + and disagreement with –."
    document.add_paragraph(s)
    
    table = document.add_table(rows=7, cols=2)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Agree/Disagree'
    
    for i in range(6):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = labels[i] 
    
    document.add_page_break()

margin = 0.5
sections = document.sections
for section in sections:
    section.top_margin = Inches(.3)
    section.bottom_margin = Inches(.2)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)
document.save('../output/patient_descriptions.docx')